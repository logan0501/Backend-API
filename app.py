from __future__ import print_function
import string
from flask import Flask,jsonify,request
import firebase_admin
from firebase_admin import credentials
from firebase_admin import firestore
import os
from dotenv import load_dotenv
import requests
import json
import os.path
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
from googleapiclient.http import MediaFileUpload
import tempfile
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
import re
from tqdm import tqdm
import docx2txt
from nltk import word_tokenize
import nltk
import heapq
import ssl
import os
from docx import Document
import shutil
import random
import nltk
from nltk.corpus import stopwords
# nltk.download('stopwords')
stopwords = stopwords.words('english')
# try:
#     _create_unverified_https_context = ssl._create_unverified_context
# except AttributeError:
#     pass
# else:
#     ssl._create_default_https_context = _create_unverified_https_context

# nltk.download('punkt')
# nltk.download('stopwords')

load_dotenv()
API_KEY = os.getenv('KEY')

cred = credentials.Certificate("secret_key.json")
firebase_admin.initialize_app(cred)
firestoreDb = firestore.client()

SCOPES = ['https://www.googleapis.com/auth/drive']


from flask_cors import CORS, cross_origin
app = Flask(__name__)
cors = CORS(app)


def get_gdrive_service():
    creds = None
    if os.path.exists('token.json'):
        creds = Credentials.from_authorized_user_file('token.json', SCOPES)

    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            flow = InstalledAppFlow.from_client_secrets_file(
                'credentials.json', SCOPES)
            creds = flow.run_local_server(port=0)
        with open('token.json', 'w') as token:
            token.write(creds.to_json())

    try:
        return build('drive', 'v3', credentials=creds)
    except HttpError as error:
        # TODO(developer) - Handle errors from drive API.
        print(f'An error occurred: {error}')
service = get_gdrive_service()

def createFolder(service,uid):
    try: 
        file_metadata = {
            'name': uid,
            'mimeType': 'application/vnd.google-apps.folder',
            "parents":["149uIJHmu2USVA0FaYkCEyTz2l3Le7AfM"]
        }

        file = service.files().create(body=file_metadata, fields='id'
                                      ).execute()
        # print(F'Folder ID: "{file.get("id")}".')
        print("[] Folder created")
        return file.get('id')

    except HttpError as error:
        return jsonify({"error":error }),400
def download_file_from_google_drive(id, destination):
    def get_confirm_token(response):
        for key, value in response.cookies.items():
            if key.startswith('download_warning'):
                return value
        return None

    def save_response_content(response, destination):
        CHUNK_SIZE = 32768
        # get the file size from Content-length response header
        file_size = int(response.headers.get("Content-Length", 0))
        # extract Content disposition from response headers
        content_disposition = response.headers.get("content-disposition")
        # parse filename
        filename = re.findall("filename=\"(.+)\"", content_disposition)[0]
        print("[+] File size:", file_size)
        print("[+] File name:", filename)
        progress = tqdm(response.iter_content(CHUNK_SIZE), f"Downloading {filename}", total=file_size, unit="Byte", unit_scale=True, unit_divisor=1024)
        with open(destination, "wb") as f:
            for chunk in progress:
                if chunk: # filter out keep-alive new chunks
                    f.write(chunk)
                    # update the progress bar
                    progress.update(len(chunk))
        progress.close()

    URL = "https://docs.google.com/uc?export=download"
    session = requests.Session()
    response = session.get(URL, params = {'id': id}, stream=True)
    print("[+] Downloading", response.url)
    token = get_confirm_token(response)
    if token:
        params = {'id': id, 'confirm':token}
        response = session.get(URL, params=params, stream=True)
    save_response_content(response, destination)  
def download(file_id):
    service = get_gdrive_service()
    filename = "files/"+file_id+".docx"
    service.permissions().create(body={"role": "reader", "type": "anyone"}, fileId=file_id).execute()
    download_file_from_google_drive(file_id, filename)



def SignUp(email,password,userData,userType):
    signUpUrl = "https://identitytoolkit.googleapis.com/v1/accounts:signUp?key="+API_KEY
    data =json.dumps({"email":email,"password":password,"returnSecureToken":True})
    res = requests.post(signUpUrl,data=data)
    res=res.json()
    if("localId" in res):
        print("[] "+userType+" created student user account")
        uid = res['localId']
        userData['uid'] = uid
        folderId = createFolder(service,uid)
        if folderId:
            userData['folderId']=folderId
            dbRes = firestoreDb.collection(userType+'s').document(uid).set(userData)
            if (dbRes):
                return jsonify({"data":{userType:res},"message":"Successfully created new user"}),200
            else:
                return jsonify({"error":"Error occured while creating folder in gdriver" }),400
        else:
            return jsonify({"error":"Error occured while uploading data to firestore" }),400
    if("error" in res):
        print("[] Error occurred")
        return jsonify({"error":res["error"]["message"] }),400

def LogIn(email,passoword,userType):
    loginUrl = "https://identitytoolkit.googleapis.com/v1/accounts:signInWithPassword?key="+API_KEY
    data =json.dumps({"email":email,"password":passoword,"returnSecureToken":True})
    res = requests.post(loginUrl,data=data)
    res = res.json()
    if("localId" in res):
        doc_ref = firestoreDb.collection(userType+'s').document(res['localId'])
        doc = doc_ref.get()
        if doc.exists:
            print("[] "+userType+" data found")
            isAdmin = False
            if 'isAdmin' in doc.to_dict():
                isAdmin=True
            return jsonify({"data":{userType:res},"isAdmin":isAdmin,"message":userType+" data found"}),200
        else:
            print("[] user trying to cheat")
            return jsonify({"error":"Access denied for the user"}),400
    if("error" in res):
        print("[] Error occured")
        return jsonify({"error":res["error"]["message"] }),400

# ---- Signup users ----
@app.route("/signup-student",methods =['POST','GET'])
def signUpStudent():
    if (request.method == 'POST'):
        studentData = request.get_json()
        studentEmail = studentData['studentEmail']
        studentPassword = studentData['studentPassword']
        print("[] Student data received for signup")
        return SignUp(studentEmail,studentPassword,studentData,'student')
        
@app.route("/signup-teacher",methods=['POST','GET'])    
def signUpTeacher():
    if (request.method =='POST'):
        teacherData = request.get_json()
        teacherEmail = teacherData['teacherEmail']
        teacherPassword = teacherData['teacherPassword']
        print("[] Teacher data received for signup")
        return SignUp(teacherEmail,teacherPassword,teacherData,'teacher')

@app.route("/login-student",methods=["POST","GET"])
def LoginStudent():
    if (request.method == 'POST'):
        studentData = request.get_json()
        studentEmail = studentData['studentEmail']
        studentPassword = studentData['studentPassword']
        return LogIn(studentEmail,studentPassword,'student')

@app.route("/login-teacher",methods=["POST","GET"])
def LoginTeacher():
    if (request.method == 'POST'):
        teacherData = request.get_json()
        teacherEmail = teacherData['teacherEmail']
        teacherPassword = teacherData['teacherPassword']
        return LogIn(teacherEmail,teacherPassword,'teacher')
     

def uploadFile(service,temp,filename,folderId,mimetype='application/pdf'):
 
    # 149uIJHmu2USVA0FaYkCEyTz2l3Le7AfM
    try:

        file_metadata = {'name': filename,"parents":[folderId]}
        media = MediaFileUpload(temp.name,
                                mimetype)
        # pylint: disable=maybe-no-member
        file = service.files().create(body=file_metadata, media_body=media,
                                      fields='webViewLink,id',supportsAllDrives=True).execute()

        print(F'[] File uploaded')
        return [file.get('id'),file.get('webViewLink')]
    except HttpError as error:
        print(F'An error occurred: {error}')
        file = None
    os.remove(temp.name)
    return [file.get('id'),file.get('webViewLink')]

def sendNoteVerificationEmail(user,file,uId,fileId,subject):
    print("[] Email service initiated")
    temp = tempfile.NamedTemporaryFile(delete=False)
    file.save(temp.name)
    EMAIL = os.getenv('EMAIL')
    PASSWORD = os.getenv('PASSWORD')
    fromaddr = EMAIL
    toaddr = user['studentEmail']
    msg = MIMEMultipart('alternative')
    msg['Subject'] = "Verification of notes"
    msg['From'] = fromaddr
    msg['To'] = toaddr
    html ="""<!DOCTYPE html>
    <html lang="en">
    <head></head>
    <body>
        <h1>Notely : Notification for verifying the notes</h1>
        <p>
        """+user['studentName']+""" has published a note on """+subject+""", we would like to know
        whether the content in the note shared by them is appropriate and relavant.
        After reviewing the attachment kindly accept or reject the notes.
        </p>
        <p>
          <a href="https://notely-cit.web.app/note-accept?voterid="""+user['uid']+"""&userid="""+uId+"""&noteid="""+fileId+"""\">Click Here</a>
          <span> to accept the note.</span>
        </p>
        <p>
          <a href="https://notely-cit.web.app/note-reject?voterid="""+user['uid']+"""&userid="""+uId+"""&noteid="""+fileId+"""\">Click Here</a>
          <span> to reject the note.</span>
        </p>
        <br/>
        <h4>Thank you</h4>
        <p>Regards,</p>
        <p>Team Notely.</p>
    </body>
    </html>

        """
    msg.attach(MIMEText(html, 'html'))
    attachment = open(temp.name, "rb")
    p = MIMEBase('application', 'octet-stream')
    p.set_payload((attachment).read())
    encoders.encode_base64(p)
    p.add_header('Content-Disposition', "attachment; filename= %s" % file.filename)
    msg.attach(p)
    s = smtplib.SMTP('smtp.gmail.com', 587)
    s.starttls()
    s.login(EMAIL,"udvgimfcrssasnlf")
    message = msg.as_string()
    s.sendmail(EMAIL, toaddr, message)
    s.quit()
    os.remove(temp.name)
    print("[] Email sent for verification")
@app.route("/upload-note",methods=["POST","GET"])
def uploadNotes():
    if (request.method=='POST'):
        file = request.files['file']
        uploaddata = json.loads(request.form['data'])
        uId = uploaddata['uId']
        userType = uploaddata["userType"]
        subject= uploaddata['subjectName']
        temp = tempfile.NamedTemporaryFile(delete=False)
        file.save(temp.name)
        doc_ref  = firestoreDb.collection(userType+'s').document(uId)

        if doc_ref:
            doc = doc_ref.get().to_dict()
            folderId = doc["folderId"]
            [fileId,fileLink] = uploadFile(service,temp,file.filename,folderId)
            if fileId:
                uploaddata["isAdminApproved"]=False
                uploaddata["fileId"]=fileId
                uploaddata["fileLink"]=fileLink
                uploaddata["upVotes"]=[]
                uploaddata["downVotes"]=[]
                firestoreDb.collection(uId).document(fileId).set({"fileId":fileId,"uId":uId})
                dbRes = firestoreDb.collection("notes").document(fileId).set(uploaddata)
                if dbRes:
                    docs = firestoreDb.collection(u'students').stream()
                    userDict =dict()
                    for doc in docs:
                        userDict[doc.id] = doc.to_dict()
                    randomUsers=[]
                    while len(randomUsers)<2:
                        choice = random.choice( list(userDict.keys()))
                        if choice not in randomUsers and choice!=uId:
                            randomUsers.append(choice)
                    # print(randomUsers)
                    for user in randomUsers:
                        sendNoteVerificationEmail(userDict[user],file,uId,fileId,subject)
                    return jsonify({"message":"File uploaded sucessfully and currently in verification phase","filedId":fileId }),200
                else:
                    return jsonify({"error":"Error while uploading the file"}),500
            else:
                print("[] Error while uploading the file")
                return jsonify({"error":"Error while uploading the file"}),500
        else:
            print("[] Error while retreiving the document from firebase.")
            return jsonify({"error":"Error while retreving the document from firebase"}),500
    

def sendNoteVerifiedEmail(recEmail,html):
    print("[]Verified Email Initiated")
    EMAIL = os.getenv('EMAIL')
    PASSWORD = os.getenv('PASSWORD')
    fromaddr = EMAIL
    toaddr = recEmail

    msg = MIMEMultipart('alternative')
    msg['Subject'] = "Verification of notes"
    msg['From'] = fromaddr
    msg['To'] = toaddr

    msg.attach(MIMEText(html, 'html'))

    s = smtplib.SMTP('smtp.gmail.com', 587)
    s.starttls()
    s.login(EMAIL, PASSWORD)
    message = msg.as_string()
    s.sendmail(EMAIL, toaddr, message)
    s.quit()
    print("[] Verified Email Sent")

@app.route("/note-accept",methods=['POST','GET'])
def noteAccept():
    if (request.method =='POST'):
        noteData = request.get_json()
        voterId = noteData['voterId']
        userId = noteData['userId']
        noteId = noteData['noteId']
        noteData = firestoreDb.collection('notes').document(noteId).get().to_dict()
        
        if  "verifiedBy" in noteData:
            if voterId in noteData["verifiedBy"]:
                return jsonify({"error":"user trying to accept again"}),400
            else:
                noteData["verifiedBy"].append(voterId)
        else:
            noteData["verifiedBy"]=[voterId]
        if len(noteData["verifiedBy"])==2:
            noteData['isAdminApproved']=True
            html = """<!DOCTYPE html>
            <html lang="en">
            <head></head>
            <body>
                <h1>Greetings from Notely</h1>
                <p>Your notes on """+noteData['subjectName']+""" got successfully verified.</p>
                <h4>Thank you</h4>
                <p>Regards,</p>
                <p>Team Notely.</p>
            </body>
            </html>
            """
            email = firestoreDb.collection('students').document(noteData['uId']).get().to_dict()['studentEmail']
            sendNoteVerifiedEmail(email,html)
        res = firestoreDb.collection('notes').document(noteId).update(noteData)
        return jsonify({"message":"Verification updated"}),200

@app.route("/note-reject",methods=['POST','GET'])
def noteReject():
    if (request.method =='POST'):
        noteData = request.get_json()
        voterId = noteData['voterId']
        userId = noteData['userId']
        noteId = noteData['noteId']
        noteData = firestoreDb.collection('notes').document(noteId).get().to_dict()
        
        if  "rejectedBy" in noteData:
            if voterId in noteData["rejectedBy"]:
                return jsonify({"error":"user trying to reject again"}),400
            else:
                noteData["rejectedBy"].append(voterId)
        else:
            noteData["rejectedBy"]=[voterId]
        if len(noteData["rejectedBy"])==2:
            noteData['isAdminApproved']=False
            html = """<!DOCTYPE html>
            <html lang="en">
            <head></head>
            <body>
                <h1>Greetings from Notely</h1>
                <p>Your notes on """+noteData['subjectName']+""" got rejected, kindly correct it and reupload it again.</p>
                <h4>Thank you</h4>
                <p>Regards,</p>
                <p>Team Notely.</p>
            </body>
            </html>
            """
            email = firestoreDb.collection('students').document(noteData['uId']).get().to_dict()['studentEmail']
            sendNoteVerifiedEmail(email,html)
        res = firestoreDb.collection('notes').document(noteId).update(noteData)
        return jsonify({"message":"Verification updated"}),200


@app.route("/get-all-notes",methods=['POST','GET'])
def getAllNotes():
    if request.method=='GET':
        docs = firestoreDb.collection('notes').where(u'isAdminApproved', u'==', True).stream()
        notesList=[]
        for doc in docs:
            notesList.append(doc.to_dict())
        return jsonify({"notes":notesList}),200
    
@app.route("/get-all-notes-admin",methods=['POST',"GET"])
def getAllNotesAdmin():
    if request.method=='GET':
        docs = firestoreDb.collection('notes').stream()
        notesList=[]
        for doc in docs:
            dictionary = doc.to_dict()
            if 'downVotes' in dictionary and 'upVotes' in dictionary:
                if len(dictionary['downVotes'])-len(dictionary['upVotes'])>=3:
                    notesList.append(doc.to_dict())
        return jsonify({"notes":notesList}),200    
    
@app.route("/update-note-admin",methods=['POST',"GET"])
def updateNoteAdmin():
    if request.method=='POST':
        noteData = request.get_json()
        type = noteData['type']
        noteId = noteData['noteId']
        if type=="accept":
            res = firestoreDb.collection('notes').document(noteId).set({
                'downVotes':[]
            },merge=True)
            if res:
                return jsonify({"message":"Note accepted successfully"}),200
            else:
                return jsonify({"message":"Note not found"}),400
        elif type=="reject":
            res = firestoreDb.collection('notes').document(noteId).delete()
            if res:
                return jsonify({"message":"Note removed from database successfully"}),200
            else:
                return jsonify({"message":"Error while deleting the note"}),400
        else:
            return jsonify({"message":"Response note found"}),400

@app.route("/update-votes",methods=['POST','GET'])
def updateNotes():
    if request.method == 'POST':
        noteData = request.get_json()
        res = firestoreDb.collection('notes').document(noteData['noteId']).set({
            "upVotes":noteData["upVotes"],
            "downVotes":noteData["downVotes"]
        },merge=True)
        print(res)
        if (res):
            return jsonify({"message":"Vote updated"}),200
        else:
            return jsonify({"message":res}),400
        
def uploadFileFromLocalFile(service,filename,folderId="1ZXN9MidQSkPIY-3OGx-nH43W77eeVfwF",faqid=''):
    file=""

    try:

        file_metadata = {'name': filename,"parents":[folderId]}
        media = MediaFileUpload("files/"+filename,
                                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        # pylint: disable=maybe-no-member
        file = service.files().create(body=file_metadata, media_body=media,
                                    fields='webViewLink,id',supportsAllDrives=True).execute()  
        print(F'[] File uploaded')
        return [file.get('id'),file.get('webViewLink')]
    
    except HttpError as error:
        print(F'An error occurred: {error}')
        file = None
    return [file.get('id'),file.get('webViewLink')]
def generate(completeQuestions):
    tokenized_words = []
    for sent in completeQuestions:
        tokenized_words.extend(word_tokenize(sent))
    stopwords = nltk.corpus.stopwords.words('english')
    word_frequencies = {}
    for word in tokenized_words:
        if word not in stopwords:
            if word not in word_frequencies.keys():
                word_frequencies[word] = 1
            else:
                word_frequencies[word] += 1
    print(word_frequencies)
    maximum_frequncy = max(word_frequencies.values())

    for word in word_frequencies.keys():
        word_frequencies[word] = (word_frequencies[word]/maximum_frequncy)

    sentence_scores = {}

    for sent in completeQuestions:
        for word in nltk.word_tokenize(sent.lower()):
            if word in word_frequencies.keys():
                if len(sent.split(' ')) < 30:
                    if sent not in sentence_scores.keys():
                        sentence_scores[sent] = word_frequencies[word]
                    else:
                        sentence_scores[sent] += word_frequencies[word]
    summary_sentences = heapq.nlargest(len(completeQuestions)//2, sentence_scores, key=sentence_scores.get)
    return summary_sentences
def getQuestionsFromText(docText):
    docText=docText.split("\n")
    twoMarks=[]
    fiveMarks=[]
    visited11marks=False
    for sentence in docText:
        newSentence = sentence.strip().lower()
        if len(newSentence)!=0:
            if (newSentence!='2 marks' and newSentence!='11 marks'):
                if not visited11marks:
                    if (newSentence.find(".",0,3)!=-1):
                        newSentence = newSentence[newSentence.find(".",0,3)+2:]
                        if newSentence[-1]==".":
                            newSentence = newSentence[:-1]
                        twoMarks.append(newSentence)
                    else:
                        if newSentence[-1]==".":
                            newSentence = newSentence[:-1]
                        twoMarks.append(newSentence)
                else:
                    if (newSentence.find(".",0,3)!=-1):
                        newSentence = newSentence[newSentence.find(".",0,3)+2:]
                        if newSentence[-1]==".":
                            newSentence = newSentence[:-1]
                        fiveMarks.append(newSentence)
                    else:
                        if newSentence[-1]==".":
                            newSentence = newSentence[:-1]
                        twoMarks.append(newSentence)
            if (newSentence=='11 marks'):
                visited11marks=True

    return twoMarks,fiveMarks

def generateFAQ(subjectName,docSubName,folderId,faqid=''):
  
    completeTwoMarkQuestions =[]
    completeFiveMarkQuestions=[]
    documentTexts=[]
    for file in os.listdir("files"):
        if file.endswith(".docx"):
            documentTexts.append(docx2txt.process("./files/"+file))
    for documentText in documentTexts:
        twoMarks,fiveMarks = getQuestionsFromText(documentText)
        completeTwoMarkQuestions.extend(twoMarks)
        completeFiveMarkQuestions.extend(fiveMarks)
    faTwoMarks = generate(completeTwoMarkQuestions)
    faFiveMarks = generate(completeFiveMarkQuestions)
    document = Document()
    document.add_heading(subjectName)
    document.add_paragraph('2 Marks')
    def arrayToString(arr):
        string=''
        for i in range(len(arr)):
            string+=str(i+1)+". "+arr[i].capitalize()+"\n"
        return string
    faTwoMarksContent=arrayToString(faTwoMarks)
    document.add_paragraph(faTwoMarksContent)
    document.add_paragraph('11 Marks')
    faFiveMarksContent=arrayToString(faFiveMarks)
    document.add_paragraph(faFiveMarksContent)
    document.save("files/"+subjectName+".docx")
    fileDetails = uploadFileFromLocalFile(service=service,filename=subjectName+".docx", faqid=faqid)
    doc_ref = firestoreDb.collection('questionPapers').document(docSubName)
    result = doc_ref.set({
    'faqId':fileDetails[0],
    'faqLink':fileDetails[1],
}, merge=True)
    return {
    'faqId':fileDetails[0],
    'faqLink':fileDetails[1],
}


@app.route("/get-faq-and-papers",methods=['POST','GET'])
def getFaq():
    if request.method == "POST":
        questionPaperData = request.get_json()
        subjectName = questionPaperData["subjectName"]
        subjectName = subjectName.replace(" ","").lower()
        docs = firestoreDb.collection('questionPapers').document(subjectName).get().to_dict()
        if docs:
            availablePapers = docs["files"]
            requestedTypePapers=[]
            for paper in availablePapers:
                for paperDetails in paper.values():
                    if paperDetails['type']==questionPaperData["type"]:
                        requestedTypePapers.append(paperDetails['link'])
            reponseData = {
                "subjectName":docs["subjectName"],
                "subjectCode":docs["subjectCode"],
                "files":requestedTypePapers
            }
            # availablePapers = docs["files"]
            # requestedPaperIds=list(availablePapers[0].keys())
            # os.mkdir("files")
            # for id in requestedPaperIds:
            #     download(id)
            # generateFAQ(docs["subjectName"],subjectName)
            if 'faqId' in docs:
                return jsonify({"message":"faq found","faqLink":docs['faqLink'],"responseData":reponseData}),200
            else:
                return jsonify({"message":"Data not found"}),400
        else:
            return jsonify({"message":"Data not found"}),400

def createFolderForFAQ(service,subjectName):
    try: 
        file_metadata = {
            'name': subjectName,
            'mimeType': 'application/vnd.google-apps.folder',
            "parents":["15kx6vmuikQSjr2Ny7_7y84CEGs3zdmlX"]
        }

        file = service.files().create(body=file_metadata, fields='id'
                                      ).execute()
        # print(F'Folder ID: "{file.get("id")}".')
        print("[] Folder created")
        return file.get('id')
    except HttpError as error:
        return jsonify({"error":error }),400
    
@app.route("/generate-faq",methods=['POST','GET'])
def generateFaqFromUserGivenPapers():
    if request.method=='POST':
        files = request.files.getlist('file')
        uploaddata = json.loads(request.form['data'])
        subjectName = uploaddata["subjectName"]
        subjectCode= uploaddata["subjectCode"]
        type = uploaddata["type"]
        folderId=''
        subjectName = subjectName.replace(" ","").lower()
        doc_ref = firestoreDb.collection('questionPapers').document(subjectName)
        if not doc_ref.get().exists:
            folderId = createFolderForFAQ(service,subjectName)
            doc_ref.set({"subjectName":uploaddata["subjectName"],"subjectCode":subjectCode,"folderId":folderId})
            newfiles={}
            for file in files:
                temp = tempfile.NamedTemporaryFile(delete=False)
                file.save(temp.name)
                id,link=uploadFile(service,temp,file.filename,folderId,'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                newfiles[id]={'type':type,'link':link}
            res = doc_ref.set({'files':[newfiles]},merge=True)
            if res:
                print("File uploaded successfully")
            else:
                print("error")
        else:
            docs = doc_ref.get().to_dict()
            folderId=docs['folderId']
            newfiles=docs['files'][0]
            for file in files:
                temp = tempfile.NamedTemporaryFile(delete=False)
                file.save(temp.name)
                id,link=uploadFile(service,temp,file.filename,folderId,'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                newfiles[id]={'type':type,'link':link}
                res = doc_ref.set({'files':[newfiles]},merge=True)
                if res:
                    print("File uploaded successfully")
                else:
                    print("error")
        docs = firestoreDb.collection('questionPapers').document(subjectName).get().to_dict()
        if docs:
            availablePapers = docs["files"]
            requestedTypePapers=[]
            for paper in availablePapers:
                for paperDetails in paper.values():
                    if paperDetails['type']==type:
                        requestedTypePapers.append(paperDetails['link'])
            availablePapers = docs["files"]
            faqId=''
            if 'faqId' in docs:
                faqId=docs['faqId']
            requestedPaperIds=list(availablePapers[0].keys())
            os.mkdir("files")
            for id in requestedPaperIds:
                download(id)
            result = generateFAQ(docs["subjectName"],subjectName,folderId ,faqId)
            reponseData = {
                "subjectName":docs["subjectName"],
                "subjectCode":docs["subjectCode"],
                "files":requestedTypePapers
            }
            shutil.rmtree('files')
            if 'faqId' in result:
                return jsonify({"message":"faq found","faqLink":result['faqLink'],"responseData":reponseData}),200
            else:
                return jsonify({"message":"Data not found"}),400
        else:
            return jsonify({"message":"Data not found"}),400
            

@app.route("/create-private-note",methods=['POST','GET'])
def createPrivateNotes():
    if request.method =='POST':
        noteData = request.get_json()
        print(noteData)
        noteContent=noteData["data"]
        noteUserId=noteData["uId"]
        doc_ref = firestoreDb.collection(noteUserId).add({"type":"private","uId":noteUserId,"data":noteContent})
        res =firestoreDb.collection(noteUserId).document(doc_ref[1].id).set({"noteId":doc_ref[1].id},merge=True)
        if res:
            return jsonify({"message":"Private note created successfully","noteId":doc_ref[1].id}),200
        else:
            return jsonify({"message":"Error while creating a private note"}),400

@app.route("/get-private-notes",methods=['POST','GET'])
def getPrivateNotes():
    if request.method=='POST':
        userData = request.get_json()
        uid = userData["uId"]
        docs = firestoreDb.collection(uid).stream()
        privateNotes=[]
        for doc in docs:
            docId=doc.id
            noteDict = doc.to_dict()
            if "type" in noteDict:
               privateNotes.append(noteDict)
        return jsonify({"message":"Private notes","notes":privateNotes}),200

@app.route("/get-all-list-of-faqs",methods=['POST','GET'])
def getAllListOfFaqs():
    if request.method=='GET':
        docs = firestoreDb.collection("questionPapers").stream()
        faqSubjectNames=[]
        for doc in docs:
            subjectDoc = doc.to_dict()
            if "faqId" in subjectDoc:
                faqSubjectNames.append(subjectDoc["subjectName"])
        return jsonify({"message":"List of available faqs","list":faqSubjectNames}),200

def generateCoursePlan(subjectName,fileId,originalSubjectName):        
    docContent = docx2txt.process("files/"+fileId+".docx")
    docContent = docContent.split("\n")
    questions=[]
    for sentence in docContent:
        newSentence = sentence.strip()
        if len(newSentence)>0 and newSentence[0].isdigit() and newSentence.lower() != '2 marks' and newSentence.lower()!='11 marks':
            newSentence = newSentence[newSentence.find(".",0,3)+2:]
            questions.append(newSentence)            
    coursePlanContent = docx2txt.process("files/"+subjectName+"CoursePlan.docx")
    coursePlanContent = coursePlanContent.split("\n")
    coursePlanTopics = []
    for topic in coursePlanContent:
        if len(topic.strip()) > 0 and "course plan" not in topic.lower():
            coursePlanTopics.append(topic)
    def clean_string(text):

        text = "".join([word for word in text if word not in string.punctuation])
        text = text.lower()
        return text
    cleanedSentences = list(map(clean_string,questions))
    cleanedPhrases = list(map(clean_string,coursePlanTopics))
    importantTopics = set()
    notFoundQuestions = []


    def check_common_phrases(sentence1, sentence2):
        # Tokenize the sentences into words or phrases
        words1 = set(sentence1.lower().split())
        words2 = set(sentence2.lower().split())
        # Find the intersection of the sets
        common_phrases = words1.intersection(words2)
        score=0
        for phrase in common_phrases:
            if phrase not in stopwords:
                score+=1  
            if score>=2:
                return [sentence2,score,sentence1]
        return [-1,-1,-1]

    coursePlan = dict()
    questions = set()

    for sentence in cleanedSentences:
        for phrase in cleanedPhrases:
            topic, score,question = check_common_phrases(sentence,phrase)
            if topic !=-1:
                questions.add(question)
                if topic in coursePlan:
                    if coursePlan[topic]<score:
                        coursePlan[topic]=score
                else:
                    coursePlan[topic]=score

    importantTopics = list(coursePlan.keys())
    importantTopics = sorted(importantTopics)
    notFoundQuestions = list(set(cleanedSentences)-questions)
    document = Document()
    document.add_heading(originalSubjectName+" important topics")
    def courseDisplayHelper(arr):
        string=''
        for i in range(len(arr)):
            string+=arr[i][0]+". "+arr[i][2:].capitalize()+"\n"
        return string
    importantTopicsText=courseDisplayHelper(list(importantTopics))
    document.add_paragraph(importantTopicsText)
    def arrayToString(arr):
        string=''
        for i in range(len(arr)):
            string+=str(i+1)+". "+arr[i].capitalize()+"\n"
        return string
    document.add_paragraph('Additional Questions not found in course plan')
    notFoundQuestionsContent=arrayToString(notFoundQuestions)
    document.add_paragraph(notFoundQuestionsContent)
    document.save("files/"+subjectName+".docx")
    fileDetails = uploadFileFromLocalFile(service,subjectName+".docx",folderId='1uUXRF5oHe6IbOSxnLy45p4yIZmvw3rqb')
    doc_ref = firestoreDb.collection('questionPapers').document(subjectName)
    res = doc_ref.set({
    'importantTopicsFileId':fileDetails[0],
    'importantTopicFileLink':fileDetails[1],
}, merge=True)
    shutil.rmtree('files')
    print(res)
    if "update_time" in res:
        print("File uploaded successfully")
        return jsonify({"message":"File uploaded successfully",'link':fileDetails[1]}),200
    else:
        return jsonify({"message":"Error while uploading file"}),400

@app.route("/generate-and-get-course-plan",methods=['POST','GET'])
def generateAndGetCoursePlan():
    if request.method=='POST':
        file = request.files['file']
        uploaddata = json.loads(request.form['data'])
        originalSubjectName = uploaddata["subjectName"]
        subjectName = originalSubjectName.replace(" ","").lower()
        doc_ref = firestoreDb.collection('questionPapers').document(subjectName)
        if not doc_ref.get().exists:
            return jsonify({"message":"Subject doesn't exist"}),400
        else:
            docDict = doc_ref.get().to_dict()
            
            if "importantTopicFileLink" in docDict:
                return jsonify({"message":"File found already",'link':docDict["importantTopicFileLink"]}),200
            else:
                doc = firestoreDb.collection('questionPapers').document(subjectName).get().to_dict()
                faqFileId = doc["faqId"]
                os.mkdir("files")
                download(faqFileId)
                file.save("files/"+subjectName+"CoursePlan.docx")
                return generateCoursePlan(subjectName,faqFileId,originalSubjectName)

# shutil.rmtree('files')
if __name__ == "__main__":
    app.run(debug=False,port=8001)


# flask run